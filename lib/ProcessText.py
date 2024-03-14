import fitz
import os, io
from thefuzz import fuzz
import pandas as pd
from docx import Document as Document_Doc

#from lib.ProcessMyPDF import *

class ProcessText:

    VALID_FORMATS_FITZ = ["pdf","xps", "epub", "mobi", "fb2", "cbz", "svg"]
    VALID_FORMATS = ["txt"]
    DOCS_FORMAT = ["docx"]
    
    @staticmethod
    def check_if_header(pages):
        if len(pages) > 1:    
            pages[0] = pages[0].strip()
            first_page = pages[0].splitlines()
            for idx, line in enumerate(first_page): 
                for s in pages[1:]:
                     s = s.strip()
                     next_page = s.splitlines()
                     #print("line: "+str(idx) + "|* page0*|: " + line + " |*page"+ "*|: "+ next_page[idx])
                     if len(next_page) <= idx or fuzz.ratio(next_page[idx],line) < 90:
                        return idx
        return 0
        
    @staticmethod    
    def check_if_footer(pages):
        if len(pages) > 1:    
            pages[0] = pages[0].strip()
            first_page = pages[0].splitlines()
            idx = 0
            for line in reversed(first_page):
                for s in pages[1:]:
                    s = s.strip()
                    next_page = s.splitlines()                
                    current_line = len(next_page) - idx - 1 
                    if current_line < 0 or fuzz.ratio(next_page[current_line],line) < 90:
                        return idx
                idx+=1
        return 0
        
    @staticmethod
    def remove_header_footer(text):
        '''Remove header and footer'''
        pages = text.split("\f")
        #pages = [x for x in pages if len(x.strip())>0] #remove white slices
        header = ProcessText.check_if_header(pages)
        footer = ProcessText.check_if_footer(pages)
        if header > 0 or footer > 0:
            text_formatted = ""
            for page in pages:            
                #text_formatted = text_formatted + remove_header_footer(page, header, footer)
                page = page.strip()
                text_splitted = page.splitlines()
                result = ""
                value = len(text_splitted)
                if footer > 0:
                    value = -1 * footer
                for line in text_splitted[header:value]:
                    text_formatted = text_formatted + line + "\r\n"                
                text_formatted = text_formatted + "\f\r\n"
            
            return text_formatted.strip()
        else:
            return text

    @staticmethod
    def readFile(bytesFile, fileType):
        txt_formatted = ""                  
        splitted_text = []
        text_pages = []
        fileType = fileType.lower() 
        #bytesFile = file.file.read()            
        if fileType in ProcessText.VALID_FORMATS_FITZ:
            plain_text = ProcessText.readFileFitz(bytesFile, fileType)
            txt_formatted = ProcessText.remove_header_footer(plain_text)
            splitted_text, text_pages = ProcessText.chunk_text(txt_formatted) #fix text 

        elif fileType in ProcessText.VALID_FORMATS:
            plain_text = bytesFile.decode("utf-8").strip()
            txt_formatted = ProcessText.remove_header_footer(plain_text)
            splitted_text, text_pages = ProcessText.chunk_text(txt_formatted) #fix text 
            
        elif fileType in ProcessText.DOCS_FORMAT:
            splitted_text = ProcessText.readFileDocx(bytesFile)
            text_pages.append(1)    #docx hasn't got pages            
            
        else:
            print("Error extension")
            #return txt_formatted
        
 
        #txt_formatted = ProcessText.remove_header_footer(plain_text)
        return splitted_text, text_pages
                    
    @staticmethod        
    def readFileFitz(bytesFile, myformat):
        #TODO: https://towardsdatascience.com/extracting-text-from-pdf-files-with-python-a-comprehensive-guide-9fc4003d517
        doc = fitz.open(stream=bytesFile, filetype=myformat)
        plain_text = ""
        for page in doc: 
            
            # tables=page.find_tables()
            # for tab in page.find_tables(): 
                # # process the content of table 'tab'
                # page.add_redact_annot(tab.bbox)  # wrap table in a redaction annotation
                
            # page.apply_redactions()  # erase all table text
            
            blocks = page.get_text("blocks",sort=True)
            for block in blocks:
                if block[6] == 0: #0 text, 1 image
                    plain_text += block[4]  #4 = text
                    item = block[4].rstrip()
                    if len(item) > 0 and item[-1] in [".","?","!"]:
                        plain_text += "\r\n"
            plain_text = plain_text + "\f\r\n"
        return plain_text.strip()
    
    @staticmethod
    def readFileDocx(bytesFile):
        source_stream = io.BytesIO(bytesFile)
        mydocument = Document_Doc(source_stream)
        source_stream.close()       
        #I = []
        #plain_text = ""
        
        #paragraphs = mydocument.paragraphs
        paragraphs = [x.text.strip() for x in mydocument.paragraphs]
        #for i in range(len(paragraphs)):
        #    plain_text += paragraphs[i].text.strip()
        #    plain_text += "\r\n"
            #if 'graphicData' in par[i]._p.xml:
            #    I.append(i)
        #print(I)
        #return (plain_text,I)
        #return plain_text.strip()
        return paragraphs
    
    #@staticmethod 
    # def extract_Docx_tables(document):            
        # tables = []
        # for table in document.tables:
            # df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
            # for i, row in enumerate(table.rows):
                # for j, cell in enumerate(row.cells):
                    # if cell.text:
                        # df[i][j] = cell.text
            # tables.append(pd.DataFrame(df))
        # return tables
    
    @staticmethod    
    def chunk_text(input):
        '''
        parse text by empty line. clear emtpy lines and whites in the end.
        '''
        text = ""
        splitted_text = []
        text_page = []
        buf = io.StringIO(input)
        file_content = buf.readlines()
        page_counter = 1
        for line in file_content:
            #line = line.decode("utf-8")       
            aux = line.strip()
            if len(aux) == 0 :  #breakline
                if len(text) > 0 and text.strip().endswith("."):
                    splitted_text.append(text.strip())                
                    text_page.append(page_counter)
                    text = ""
                if "\f" in line:    #end of page
                    page_counter+=1
                #continue
            else:
                if not aux.endswith(".") and not aux.endswith(":"):
                    line = aux + " "
                text = text + line

        return splitted_text, text_page